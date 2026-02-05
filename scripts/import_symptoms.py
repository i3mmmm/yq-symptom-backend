#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
医疗症状收集与分析系统 - 症状数据导入脚本

功能：
1. 解析DOCX文档中的300个症状信息
2. 按区域划分导入到symptoms表
3. 提取症状原因中的标签信息，建立symptom_tags关联
4. 确保数据导入完整，无缺失症状

使用SQLAlchemy核心直接操作数据库，避免Flask依赖问题。
"""

import os
import sys
import re
import docx
from pathlib import Path

# 添加项目路径
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '../app'))

# 使用SQLAlchemy核心
from sqlalchemy import create_engine, MetaData, Table, Column, Integer, String, Text, DateTime, ForeignKey
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
from datetime import datetime

# 标签定义：从症状原因中提取的关键词
TAG_KEYWORDS = [
    '营养', '排毒', '毒素', '习惯', '体质', '免疫力', 
    '微循环', '循环', '内分泌', '情绪', '寒湿', '温度'
]

# 区域划分
AREA_DEFINITIONS = {
    'red': {'name': '红色区域', 'range': (1, 50), 'description': '心、小肠'},
    'green': {'name': '绿色区域', 'range': (51, 100), 'description': '肝、胆'},
    'white': {'name': '白色区域', 'range': (101, 150), 'description': '肺、大肠'},
    'black': {'name': '黑色区域', 'range': (151, 200), 'description': '肾、膀胱'},
    'yellow': {'name': '黄色区域', 'range': (201, 250), 'description': '脾、胃'},
    'blue': {'name': '蓝色区域', 'range': (251, 300), 'description': '妇科病'}
}

# 数据库配置
DATABASE_PATH = 'health_system_dev.db'

# 创建数据库引擎
engine = create_engine(f'sqlite:///{DATABASE_PATH}', echo=False)
Base = declarative_base()

# 定义模型（与app/models.py保持一致）
class Symptom(Base):
    """症状主表模型"""
    __tablename__ = 'symptoms'
    
    id = Column(Integer, primary_key=True)  # 1-300，与文档序号一致
    name = Column(String(50), nullable=False)
    area = Column(String(20), nullable=False)
    description = Column(Text, nullable=False)
    precautions = Column(Text, nullable=True)
    contraindications = Column(Text, nullable=True)
    created_at = Column(DateTime, nullable=False, default=datetime.utcnow)
    updated_at = Column(DateTime, nullable=False, default=datetime.utcnow, onupdate=datetime.utcnow)

class SymptomTag(Base):
    """症状原因标签表模型"""
    __tablename__ = 'symptom_tags'
    
    symptom_id = Column(Integer, ForeignKey('symptoms.id', ondelete='CASCADE'), primary_key=True)
    tag = Column(String(20), primary_key=True)
    created_at = Column(DateTime, nullable=False, default=datetime.utcnow)

# 创建表
Base.metadata.create_all(engine)

# 创建会话
Session = sessionmaker(bind=engine)
session = Session()

def extract_tags_from_causes(causes_text):
    """从症状原因文本中提取标签"""
    tags = set()
    
    # 查找关键词
    for keyword in TAG_KEYWORDS:
        if keyword in causes_text:
            tags.add(keyword)
    
    # 检查特殊组合
    if '益生菌' in causes_text:
        tags.add('排毒')
    
    # 确保至少有1个标签
    if not tags:
        # 如果没有找到明确的标签，添加默认标签
        tags.add('其他')
    
    return list(tags)

def parse_table1(table):
    """解析表格1（有表头）"""
    symptoms = []
    
    # 从第1行开始（跳过表头）
    for row_idx in range(1, len(table.rows)):
        row = table.rows[row_idx]
        cells = row.cells
        
        if len(cells) < 3:
            continue
            
        # 第一列：ID
        id_text = cells[0].text.strip()
        if not id_text:
            continue
            
        try:
            symptom_id = int(id_text)
        except:
            continue
        
        # 第二列：症状名称
        name = cells[1].text.strip()
        
        # 第三列：原因
        causes = cells[2].text.strip() if len(cells) > 2 else ''
        
        # 确定区域（红色）
        area = 'red'
        
        tags = extract_tags_from_causes(causes)
        
        symptom_data = {
            'id': symptom_id,
            'name': name,
            'area': area,
            'description': causes,
            'tags': tags
        }
        
        symptoms.append(symptom_data)
    
    return symptoms

def parse_table_generic(table, expected_area):
    """解析通用表格（无表头，第一行是数据）"""
    symptoms = []
    
    # 每行数据
    for row_idx in range(0, len(table.rows)):
        row = table.rows[row_idx]
        cells = row.cells
        
        # 至少需要2列（ID和名称）
        if len(cells) < 2:
            continue
            
        # 第一列：ID
        id_text = cells[0].text.strip()
        if not id_text:
            continue
            
        # 提取数字（可能包含中文标点）
        match = re.search(r'(\d+)', id_text)
        if not match:
            continue
            
        symptom_id = int(match.group(1))
        
        # 第二列：症状名称
        name = cells[1].text.strip()
        
        # 第三列：原因（如果有）
        causes = ''
        if len(cells) >= 3:
            causes = cells[2].text.strip()
        
        tags = extract_tags_from_causes(causes)
        
        symptom_data = {
            'id': symptom_id,
            'name': name,
            'area': expected_area,
            'description': causes,
            'tags': tags
        }
        
        symptoms.append(symptom_data)
    
    return symptoms

def parse_docx_file(docx_path):
    """解析DOCX文件中的所有症状"""
    print(f"正在解析文档: {docx_path}")
    
    if not os.path.exists(docx_path):
        print(f"错误：文件不存在: {docx_path}")
        return []
    
    doc = docx.Document(docx_path)
    
    print(f"文档包含 {len(doc.tables)} 个表格")
    
    all_symptoms = []
    
    # 区域顺序
    areas = ['red', 'green', 'white', 'black', 'yellow', 'blue']
    
    for table_idx, table in enumerate(doc.tables):
        if table_idx >= len(areas):
            break
            
        area = areas[table_idx]
        
        print(f"\n=== 解析 {AREA_DEFINITIONS[area]['name']} (表格 {table_idx+1}) ===")
        
        if table_idx == 0:
            # 表格1有表头
            symptoms = parse_table1(table)
        else:
            # 其他表格无表头
            symptoms = parse_table_generic(table, area)
        
        print(f"  找到 {len(symptoms)} 个症状")
        
        # 显示前几个
        for s in symptoms[:3]:
            print(f"    ID={s['id']}: {s['name']}")
        
        all_symptoms.extend(symptoms)
    
    # 排序
    all_symptoms.sort(key=lambda x: x['id'])
    
    print(f"\n=== 解析完成 ===")
    print(f"总共解析到 {len(all_symptoms)} 个症状")
    
    # 检查缺失
    symptom_ids = [s['id'] for s in all_symptoms]
    missing_ids = []
    
    for expected_id in range(1, 301):
        if expected_id not in symptom_ids:
            missing_ids.append(expected_id)
    
    if missing_ids:
        print(f"警告：缺失 {len(missing_ids)} 个症状: {missing_ids[:10]}{'...' if len(missing_ids) > 10 else ''}")
        
        # 尝试查找重复或错误ID
        id_counts = {}
        for s in all_symptoms:
            id_counts[s['id']] = id_counts.get(s['id'], 0) + 1
        
        duplicates = [id for id, count in id_counts.items() if count > 1]
        if duplicates:
            print(f"重复ID: {duplicates}")
    else:
        print("完美！所有300个症状都已解析到。")
    
    # 显示区域统计
    area_counts = {}
    for s in all_symptoms:
        area_counts[s['area']] = area_counts.get(s['area'], 0) + 1
    
    print(f"\n区域统计:")
    for area_code in areas:
        if area_code in area_counts:
            print(f"  {AREA_DEFINITIONS[area_code]['name']}: {area_counts[area_code]} 个症状")
    
    # 显示标签统计
    tag_counts = {}
    for s in all_symptoms:
        for tag in s['tags']:
            tag_counts[tag] = tag_counts.get(tag, 0) + 1
    
    print(f"\n标签统计:")
    for tag, count in sorted(tag_counts.items(), key=lambda x: x[1], reverse=True):
        print(f"  {tag}: {count}个症状")
    
    return all_symptoms

def import_to_database(symptoms_data):
    """将症状数据导入数据库"""
    print("\n=== 导入症状数据到数据库 ===")
    
    try:
        # 清空现有症状数据
        print("清空现有症状数据...")
        session.query(SymptomTag).delete()
        session.query(Symptom).delete()
        session.commit()
        
        imported_count = 0
        tag_count = 0
        
        for symptom_data in symptoms_data:
            # 创建症状记录
            symptom = Symptom(
                id=symptom_data['id'],
                name=symptom_data['name'],
                area=symptom_data['area'],
                description=symptom_data['description'],
                precautions=None,  # 可后续补充
                contraindications=None  # 可后续补充
            )
            
            session.add(symptom)
            
            # 创建标签记录
            for tag_name in symptom_data['tags']:
                symptom_tag = SymptomTag(
                    symptom_id=symptom_data['id'],
                    tag=tag_name
                )
                session.add(symptom_tag)
                tag_count += 1
            
            imported_count += 1
            
            # 每50个症状提交一次
            if imported_count % 50 == 0:
                session.commit()
                print(f"已导入 {imported_count} 个症状...")
        
        # 最终提交
        session.commit()
        
        print(f"导入完成！共导入 {imported_count} 个症状，{tag_count} 个标签关系。")
        
        # 验证导入
        db_symptom_count = session.query(Symptom).count()
        db_tag_count = session.query(SymptomTag).count()
        
        print(f"\n数据库验证:")
        print(f"  Symptom表: {db_symptom_count} 条记录")
        print(f"  SymptomTag表: {db_tag_count} 条记录")
        
        if db_symptom_count == len(symptoms_data):
            print("  ✓ 症状数量验证通过！")
        else:
            print(f"  ✗ 警告：期望 {len(symptoms_data)} 个症状，实际 {db_symptom_count} 个")
        
        # 显示按区域统计
        print(f"\n区域分布:")
        for area_code in AREA_DEFINITIONS.keys():
            count = session.query(Symptom).filter_by(area=area_code).count()
            print(f"  {AREA_DEFINITIONS[area_code]['name']}: {count} 个症状")
        
        return True
        
    except Exception as e:
        print(f"导入过程中发生错误: {e}")
        import traceback
        traceback.print_exc()
        session.rollback()
        return False

def main():
    """主函数"""
    print("=== 医疗症状收集与分析系统 - 症状数据导入 ===\n")
    
    # 配置文件
    docx_path = '2021自检表300症状分析.docx'
    
    # 检查文档文件
    if not os.path.exists(docx_path):
        # 尝试从上级目录查找
        docx_path = os.path.join('..', docx_path)
        if not os.path.exists(docx_path):
            print(f"错误：找不到症状文档 '{docx_path}'")
            print("请确保 '2021自检表300症状分析.docx' 文件在当前目录或上级目录。")
            sys.exit(1)
    
    try:
        # 1. 解析文档
        symptoms_data = parse_docx_file(docx_path)
        
        if not symptoms_data:
            print("错误：未解析到任何症状数据")
            sys.exit(1)
        
        # 2. 导入数据库
        success = import_to_database(symptoms_data)
        
        if success:
            print("\n=== 数据导入成功完成 ===")
            print(f"数据库文件: {DATABASE_PATH}")
        else:
            print("\n=== 数据导入失败 ===")
            sys.exit(1)
        
    except Exception as e:
        print(f"导入过程中发生错误: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
    finally:
        session.close()

if __name__ == '__main__':
    main()